"""Executive desk report as Word (.docx) from Markdown-shaped text (python-docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from report.executive_markdown_render import (
    LineKind,
    append_markdown_runs_to_paragraph,
    parse_executive_line,
)


def write_executive_docx(markdown_text: str, path: Path) -> None:
    """Write a simple Word document from executive Markdown (headings + bullets)."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    heading_levels = {
        LineKind.H1: 1,
        LineKind.H2: 2,
        LineKind.H3: 3,
        LineKind.H4: 4,
        LineKind.H5: 5,
    }

    for raw_line in markdown_text.splitlines():
        parsed = parse_executive_line(raw_line)
        if parsed.kind == LineKind.BLANK:
            continue
        if parsed.kind in heading_levels:
            heading = doc.add_heading("", level=heading_levels[parsed.kind])
            append_markdown_runs_to_paragraph(heading, parsed.text)
            continue
        if parsed.kind == LineKind.BULLET:
            paragraph = doc.add_paragraph(style="List Bullet")
            append_markdown_runs_to_paragraph(paragraph, parsed.text)
            continue
        paragraph = doc.add_paragraph()
        append_markdown_runs_to_paragraph(paragraph, parsed.text)

    doc.save(str(path))
